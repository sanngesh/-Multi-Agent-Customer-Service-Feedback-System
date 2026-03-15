"""
Agent 4: Document Compiler
Deterministic agent that assembles verified feedback into a professionally
formatted Word document (.docx) using python-docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING


class DocumentCompiler:
    """Builds a formatted Word document from structured feedback data."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure default document styles for 1.5 spacing and professional fonts."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(4)
        paragraph_format.space_before = Pt(2)

        # Adjust margins to help keep under 5 pages
        for section in self.doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

    def _add_title(self):
        """Add the document title 'Case Feedback'."""
        title_para = self.doc.add_paragraph()
        title_para.alignment = 1  # Center
        title_run = title_para.add_run("Case Feedback")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = "Calibri"
        # Apply 1.5 spacing to title too
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        title_para.paragraph_format.space_after = Pt(12)

    def _add_case_heading(self, case_name: str):
        """Add a bold case heading (e.g., 'Case One')."""
        heading_para = self.doc.add_paragraph()
        heading_run = heading_para.add_run(case_name)
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.name = "Calibri"
        heading_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        heading_para.paragraph_format.space_before = Pt(10)
        heading_para.paragraph_format.space_after = Pt(6)

    def _add_feedback_item(self, item: dict, index: int):
        """Add a single feedback item as a numbered/bulleted list entry."""
        # Main bullet with the number
        bullet_para = self.doc.add_paragraph(style="List Bullet")
        bullet_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        bullet_para.paragraph_format.space_after = Pt(2)

        # Original Statement
        orig_label = bullet_para.add_run("Original Statement: ")
        orig_label.bold = True
        orig_label.font.size = Pt(11)
        orig_label.font.name = "Calibri"
        orig_text = bullet_para.add_run(f"\"{item.get('original', '')}\"")
        orig_text.font.size = Pt(11)
        orig_text.font.name = "Calibri"

        # Explanation
        expl_para = self.doc.add_paragraph()
        expl_para.paragraph_format.left_indent = Inches(0.5)
        expl_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        expl_para.paragraph_format.space_after = Pt(2)
        expl_label = expl_para.add_run("Explanation: ")
        expl_label.bold = True
        expl_label.font.size = Pt(11)
        expl_label.font.name = "Calibri"
        expl_text = expl_para.add_run(item.get("explanation", ""))
        expl_text.font.size = Pt(11)
        expl_text.font.name = "Calibri"

        # Alternative
        alt_para = self.doc.add_paragraph()
        alt_para.paragraph_format.left_indent = Inches(0.5)
        alt_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        alt_para.paragraph_format.space_after = Pt(8)
        alt_label = alt_para.add_run("Alternative: ")
        alt_label.bold = True
        alt_label.font.size = Pt(11)
        alt_label.font.name = "Calibri"
        alt_text = alt_para.add_run(f"\"{item.get('alternative', '')}\"")
        alt_text.font.size = Pt(11)
        alt_text.font.name = "Calibri"

    def compile(self, all_feedback: dict, output_path: str):
        """
        Compile all feedback into a Word document.

        Args:
            all_feedback: dict with keys 'Case One', 'Case Two', 'Case Three',
                         each containing a list of feedback_items.
            output_path: Path to save the .docx file.
        """
        self._add_title()

        for case_name in ["Case One", "Case Two", "Case Three"]:
            self._add_case_heading(case_name)

            feedback_items = all_feedback.get(case_name, [])
            for i, item in enumerate(feedback_items):
                self._add_feedback_item(item, i)

        # Ensure all paragraphs have 1.5 spacing (safety pass)
        for paragraph in self.doc.paragraphs:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        self.doc.save(output_path)
        return output_path
